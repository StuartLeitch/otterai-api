import json
import requests
import datetime
from otterai import OtterAI
import os
from docx import Document as DocumentDocx
from dotenv import load_dotenv
from llama_index.core import SimpleDirectoryReader
from dotenv import load_dotenv
from llama_index.embeddings.openai import OpenAIEmbedding
from llama_index.core import Settings
from llama_index.core import VectorStoreIndex, StorageContext
from llama_index.vector_stores.milvus import MilvusVectorStore
from llama_index.core.schema import Document
 

load_dotenv()
windows_path = os.getenv('WINDOWS_EXPORT_PATHNAME_FULL')

def docx_to_text(docx_path, txt_path):
    doc = DocumentDocx(docx_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(full_text))


def get_latest_load(folder_path):
    # List all files in the folder
    files = os.listdir(folder_path)
    
    # Filter out directories
    files = [f for f in files if os.path.isfile(os.path.join(folder_path, f))]
    
    # Sort files alphabetically
    files.sort(reverse=True)
    
    # Return the latest filename
    return datetime.datetime.strptime(files[0][:8], "%Y%m%d") if files else None

def extract_outline(speechObject):
    speech_outline = speechObject['data']['speech']['speech_outline']
    outline = []
    if speech_outline is None:
        return ''
    
    for item in speech_outline:
        outline.append(item['text'])
        if item['segments']:
            for segment in item['segments']:
                outline.append(f"  • {segment['text']}")
                if segment['segments']:
                    for subsegment in segment['segments']:
                        outline.append(f"    - {subsegment['text']}")
    return '\n'.join(outline)

def extract_speakers(speechObject):

    # Extract unique speaker names
    unique_speakers = set(speaker['speaker_name'] for speaker in speechObject['data']['speech']['speakers'])

    # Remove any numeric speaker names (like "14432233173")
    unique_speakers = {name for name in unique_speakers if not name.isdigit()}

    # Sort the names alphabetically
    sorted_speakers = sorted(unique_speakers)

    # Join the names into a single string
    result = ", ".join(sorted_speakers)

    return result

def extract_speakers(speechObject):

    # Extract unique speaker names
    unique_speakers = set(speaker['speaker_name'] for speaker in speechObject['data']['speech']['speakers'])

    # Remove any numeric speaker names (like "14432233173")
    unique_speakers = {name for name in unique_speakers if not name.isdigit()}

    # Sort the names alphabetically
    sorted_speakers = sorted(unique_speakers)

    # Join the names into a single string
    result = ", ".join(sorted_speakers)

    return result

def extract_date(speechObject):

    # Extract unique speaker names
    unique_speakers = set(speaker['speaker_name'] for speaker in speechObject['data']['speech']['speakers'])

    # Remove any numeric speaker names (like "14432233173")
    unique_speakers = {name for name in unique_speakers if not name.isdigit()}

    # Sort the names alphabetically
    sorted_speakers = sorted(unique_speakers)

    # Join the names into a single string
    result = ", ".join(sorted_speakers)

    return result

def extract_title(speechObject):

    # Extract unique speaker names
    unique_speakers = set(speaker['speaker_name'] for speaker in speechObject['data']['speech']['speakers'])

    # Remove any numeric speaker names (like "14432233173")
    unique_speakers = {name for name in unique_speakers if not name.isdigit()}

    # Sort the names alphabetically
    sorted_speakers = sorted(unique_speakers)

    # Join the names into a single string
    result = ", ".join(sorted_speakers)

    return result

PAGE_SIZE = 100
DAYS_AHEAD = 15


while True:

    latest_load_date = get_latest_load(windows_path)
    print(f'Top file: {latest_load_date}')

    if latest_load_date is None:
        # set to 5 years ago
        latest_load_date = datetime.datetime.now() - datetime.timedelta(days=365*1)

    otter = OtterAI()
    otter.login(os.getenv('OTTER_USER'), os.getenv('OTTER_PASS'))

    def calc_dt_offset(top_file):
        offset = top_file + datetime.timedelta(days=DAYS_AHEAD)
        return str(int(offset.timestamp()))

    offset_date = calc_dt_offset(latest_load_date)
    speechesObject = otter.get_speeches(page_size=PAGE_SIZE, last_load_ts=offset_date if latest_load_date else None)

    speeches = speechesObject['data']['speeches']

    print(f'Speeches found: {len(speeches)}')

    # print(otter.get_user())

    files = os.listdir(windows_path)

    found_file_to_download = False
    skipped_downloads = False

    for speech in reversed(speeches):
        # print(speech)
        if speech['summary'] is not None and speech['summary'] != '':
            speech_id = speech['otid']

            if speech['duration'] < 120:
                print('Speech too short, skipping ')
                continue

            formatted_date = datetime.datetime.fromtimestamp(speech['start_time']).strftime("%Y%m%d %H:%M:%S")
            save_name = formatted_date + " " + (speech['title'] if not speech['title']==None else speech_id).replace('/', '-')
    
            # print (formatted_date + '  ' + speech['title'] if not speech['title'] == None else '')
            # continue

            abstract = otter.get_abstract(speech_id=speech_id)

            if save_name + '.txt' in files:
                print('File already exists, skipping ' + save_name + '.txt')
                skipped_downloads = True
                continue

            speechObject = otter.get_speech(speech_id=speech_id)

            # # save speech in JSON to file
            # with open(windows_path + save_name + 'TEST.txt', 'w', encoding='utf-8') as file:
            #     json.dump(speech, file, indent=4, ensure_ascii=False)

            title = '<Title>' + speechObject['data']['speech']['title'] if not speechObject['data']['speech']['title'] is None else '' + '</Title>'
            date_time = '<DateTime>' + formatted_date + '</DateTime>'
            outline = '<Outline>' + '\n' + extract_outline(speechObject) + '\n</Outline>'
            speakers = '<Speakers>' + '\n' + extract_speakers(speechObject) + '\n</Speakers>'

            header = title + '\n' + date_time + '\n' + speakers + '\n' + abstract + '\n' + outline 

            # print(speechObject)

            # print(otter.get_speakers())

            # print(otter.query_speech(speech_id=speech_id, query='What was said in this meeting?'))
            print('Downloading ' + save_name)
            found_file_to_download = True
            downloadedspeech = otter.download_speech(speech_id=speech_id, name=save_name, pathname=windows_path, fileformat='txt')

            download_file_name_full = windows_path + save_name + '.txt'

            # docx_to_text(windows_path + save_name + '.docx', download_file_name_full)
            # # Delete the docx file
            # os.remove(windows_path + save_name + '.docx')

            # Read  newly created text file 
            with open(download_file_name_full, 'r') as file:
                transcript_text = file.readlines()
                transcript_text.insert(0, '\n\n' + header + '\n\n' + 'Transcript begins here:')

            # save speech to file
            with open(download_file_name_full, 'w', encoding='utf-8') as file:
                file.write('\n'.join(transcript_text))


            # Create a Document object
            document_full_transcript = Document(
                text='\n'.join(transcript_text),
                metadata={
                    "filename": download_file_name_full,
                     "date": formatted_date,
                     "speakers": [item['speaker_name'] for item in speechObject['data']['speech']['speakers']],
                     
                     }
            )

            embed_model = OpenAIEmbedding(model="text-embedding-3-large", embed_batch_size=10)

            Settings.embed_model = embed_model

            vector_store_full_transcripts = MilvusVectorStore(
                uri=os.getenv('MILVUS_SERVER'), token=os.getenv('MILVUS_API_KEY'), dim=3072, overwrite=False, collection_name=os.getenv('MILVUS_COLLECTION_NAME_FULL_TRANSCRIPTS')
            )
            storage_context_full_transcripts = StorageContext.from_defaults(vector_store=vector_store_full_transcripts)
            index_full_transcripts = VectorStoreIndex.from_documents(documents=[document_full_transcript], storage_context=storage_context_full_transcripts)

            # Create a Document object
            document_summary_transcript = Document(
                text=header,
                metadata={
                    "filename": download_file_name_full,
                     "date": formatted_date,
                     "speakers": [item['speaker_name'] for item in speechObject['data']['speech']['speakers']],
                     
                     }
            )


            vector_store_summary_transcripts = MilvusVectorStore(
                uri=os.getenv('MILVUS_SERVER'), token=os.getenv('MILVUS_API_KEY'), dim=3072, overwrite=False, collection_name=os.getenv('MILVUS_COLLECTION_NAME_SUMMARY_TRANSCRIPTS')
            )
            storage_context_summary_transcripts = StorageContext.from_defaults(vector_store=vector_store_summary_transcripts)
            index_summary_transcripts = VectorStoreIndex.from_documents(documents=[document_summary_transcript], storage_context=storage_context_summary_transcripts)


    if not skipped_downloads:
        print('WARNING. No files skipped. INVESTIGATE!!! Exiting....')
        break

    if not found_file_to_download:
        print('No new files to download. Exiting....')
        break

